"""
Document ingestion layer.

Loads local documents (TXT, PDF, DOCX) from app/knowledge/ (this
module's own directory), extracts their text, and returns it normalized
and ready for downstream chunking and embedding. Entirely offline:
every extraction runs against the local filesystem using local,
pure-Python libraries (pypdf, python-docx) -- no network calls, no
cloud document-processing APIs.

Modular by design: each file type is handled by one small extractor
function registered in `DocumentLoader._EXTRACTORS`. Adding a new
format later (PPTX, HTML, an enterprise-specific export, etc.) means
registering one more function here -- `load_file`/`load_directory`
never need to change.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from docx import Document as DocxDocument
from pypdf import PdfReader

# Documents live under app/knowledge/, not alongside this module -- no
# separate data/ folder is introduced.
DEFAULT_DOCUMENTS_DIR = Path(__file__).resolve().parent.parent / "knowledge"


class DocumentLoadError(Exception):
    """Raised when a document's text cannot be extracted.

    Covers both unsupported file types and failures while parsing an
    otherwise-supported file (e.g. a corrupted or encrypted PDF).
    """


@dataclass
class LoadedDocument:
    """A single document's extracted, normalized text plus metadata.

    Attributes:
        path: The source file's path.
        text: The extracted, normalized text content.
        metadata: Extra information about the document (extension,
            file size, character count, ...). Intentionally a plain
            dict so future enterprise-document needs (page counts,
            authors, classification tags, etc.) can be added without
            changing this dataclass's shape.
    """

    path: Path
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


def _extract_txt_text(path: Path) -> str:
    """Extract text from a plain-text file.

    Args:
        path: Path to the .txt file.

    Returns:
        The file's decoded text content. Uses a lenient decode (invalid
        byte sequences are replaced rather than raising) since plain-text
        exports from various enterprise tools aren't always strict UTF-8.
    """
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF file, concatenating all pages.

    Args:
        path: Path to the .pdf file.

    Returns:
        The concatenated text of every page, in order. Pages that yield
        no extractable text (e.g. scanned image-only pages, which would
        need OCR -- out of scope here) simply contribute nothing.
    """
    reader = PdfReader(str(path))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages_text)


def _extract_docx_text(path: Path) -> str:
    """Extract text from a Word document, including table content.

    Args:
        path: Path to the .docx file.

    Returns:
        The concatenated text of all paragraphs and table cells, in
        document order (tables rendered as "cell | cell | cell" rows).
    """
    document = DocxDocument(str(path))

    parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)


def _normalize_text(text: str) -> str:
    """Normalize extracted text for consistent downstream processing.

    Normalizes line endings, trims trailing whitespace per line,
    collapses runs of 3+ blank lines down to a single blank line, and
    strips leading/trailing whitespace overall.

    Args:
        text: The raw extracted text.

    Returns:
        The normalized text.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    normalized = "\n".join(lines)

    while "\n\n\n" in normalized:
        normalized = normalized.replace("\n\n\n", "\n\n")

    return normalized.strip()


class DocumentLoader:
    """Loads and normalizes text from local documents (TXT, PDF, DOCX).

    Fully offline, with a pluggable per-extension extractor registry so
    new formats can be added without touching `load_file` or
    `load_directory`.
    """

    _EXTRACTORS: Dict[str, Callable[[Path], str]] = {
        ".txt": _extract_txt_text,
        ".pdf": _extract_pdf_text,
        ".docx": _extract_docx_text,
    }

    @classmethod
    def supported_extensions(cls) -> List[str]:
        """List the file extensions this loader currently supports.

        Returns:
            A sorted list of supported extensions, e.g. [".docx",
            ".pdf", ".txt"].
        """
        return sorted(cls._EXTRACTORS.keys())

    @classmethod
    def is_supported(cls, path: Path) -> bool:
        """Check whether a file's extension is supported.

        Args:
            path: The file to check.

        Returns:
            True if this loader has a registered extractor for the
            file's extension, False otherwise.
        """
        return path.suffix.lower() in cls._EXTRACTORS

    def load_file(self, path: Path) -> LoadedDocument:
        """Load a single document and return its normalized text.

        Args:
            path: Path to the document (.txt, .pdf, or .docx).

        Returns:
            A LoadedDocument with the extracted, normalized text and
            basic metadata.

        Raises:
            FileNotFoundError: If `path` does not exist or is not a
                file.
            DocumentLoadError: If the file's extension is unsupported,
                or text extraction fails for any reason (e.g. a
                corrupted or encrypted file).
        """
        path = Path(path)
        if not path.is_file():
            raise FileNotFoundError(f"Document not found: {path}")

        extractor = self._EXTRACTORS.get(path.suffix.lower())
        if extractor is None:
            raise DocumentLoadError(
                f"Unsupported document type '{path.suffix}' for {path}. "
                f"Supported types: {', '.join(self.supported_extensions())}"
            )

        try:
            raw_text = extractor(path)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to extract text from {path}") from exc

        normalized = _normalize_text(raw_text)

        return LoadedDocument(
            path=path,
            text=normalized,
            metadata={
                "extension": path.suffix.lower(),
                "size_bytes": path.stat().st_size,
                "char_count": len(normalized),
            },
        )

    def load_directory(
        self, directory: Optional[Path] = None
    ) -> List[LoadedDocument]:
        """Load every supported document under a directory (recursively).

        Files with an unsupported extension are skipped. Files that are
        supported but fail to parse (e.g. a corrupted PDF) are also
        skipped, so one bad file doesn't stop the rest of the batch --
        use `load_file` directly if you need a hard failure on error.

        Args:
            directory: The directory to scan. Defaults to
                app/knowledge/ (this module's own directory) if not
                provided.

        Returns:
            A list of LoadedDocument, one per successfully loaded file,
            in sorted path order. Returns an empty list if the
            directory doesn't exist.
        """
        target_dir = Path(directory) if directory is not None else DEFAULT_DOCUMENTS_DIR
        if not target_dir.is_dir():
            return []

        documents: List[LoadedDocument] = []
        for path in sorted(target_dir.rglob("*")):
            if not path.is_file() or not self.is_supported(path):
                continue
            try:
                documents.append(self.load_file(path))
            except DocumentLoadError:
                continue

        return documents
